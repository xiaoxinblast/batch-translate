"""DOCX parser with stable position IDs and basic run-format preservation."""

import json
import re
import sys
from pathlib import Path
from typing import Iterator, Optional

try:
    from docx import Document
except ImportError:
    print("错误: 需要 python-docx，请执行 pip install python-docx")
    sys.exit(1)


def parse(filepath: Path, **opts) -> dict:
    """Extract body, table, header, and footer paragraphs from a DOCX."""
    doc = Document(str(filepath))
    entries = []
    for entry_id, paragraph, context in _iter_paragraphs(doc, filepath.name):
        text = paragraph.text.strip()
        if not text:
            continue
        source, format_info = _paragraph_source(paragraph, entry_id, text)
        entries.append({
            "id": entry_id,
            "source": source,
            "context": context,
            "note": "; ".join(format_info),
        })

    warnings = []
    text_box_count = len(doc.element.xpath(".//w:txbxContent"))
    if text_box_count:
        warnings.append(
            f"检测到 {text_box_count} 个文本框；当前版本不提取文本框内容"
        )

    result = {"source_file": filepath.name, "entries": entries}
    if warnings:
        result["warnings"] = warnings
    return result


def write(
    original_path: Path,
    translations_json: str | Path,
    output_path: Optional[Path] = None,
) -> Path:
    """Write submitted targets to their stable paragraph locations."""
    with open(translations_json, "r", encoding="utf-8") as f:
        translations = json.load(f)

    target_map = {}
    if isinstance(translations, list):
        target_map = {
            str(item["id"]): item["target"]
            for item in translations
            if isinstance(item, dict)
            and "id" in item
            and isinstance(item.get("target"), str)
        }
    elif isinstance(translations, dict) and isinstance(
        translations.get("entries"), list
    ):
        target_map = {
            str(item["id"]): item["target"]
            for item in translations["entries"]
            if isinstance(item, dict)
            and "id" in item
            and isinstance(item.get("target"), str)
        }

    doc = Document(str(original_path))
    for entry_id, paragraph, _context in _iter_paragraphs(
        doc, original_path.name
    ):
        if entry_id in target_map:
            _write_formatted_paragraph(paragraph, target_map[entry_id])

    if output_path is None:
        output_path = original_path.with_stem(original_path.stem + "_translated")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _iter_paragraphs(doc, filename: str) -> Iterator[tuple[str, object, str]]:
    heading_stack = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and paragraph.style and paragraph.style.name:
            match = re.fullmatch(r"Heading\s+(\d+)", paragraph.style.name)
            if match:
                level = int(match.group(1))
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(text)
        context = " > ".join(heading_stack) if heading_stack else filename
        yield f"body.p{index}", paragraph, context

    for table_index, table in enumerate(doc.tables):
        yield from _iter_table_paragraphs(
            table, f"body.t{table_index}", f"Table:{filename}"
        )

    seen_parts = set()
    for section_index, section in enumerate(doc.sections):
        references = (
            ("header", "headerReference_lst"),
            ("footer", "footerReference_lst"),
        )
        for story_kind, reference_attr in references:
            for reference in getattr(section._sectPr, reference_attr, []):
                variant = reference.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                ) or "default"
                story = _story_for(section, story_kind, variant)
                part_name = str(story.part.partname)
                if part_name in seen_parts:
                    continue
                seen_parts.add(part_name)
                part_id = Path(part_name).stem
                prefix = f"{story_kind}.{part_id}"
                context = f"{story_kind.title()}:{filename}:Section{section_index + 1}"
                for paragraph_index, paragraph in enumerate(story.paragraphs):
                    yield f"{prefix}.p{paragraph_index}", paragraph, context
                for table_index, table in enumerate(story.tables):
                    yield from _iter_table_paragraphs(
                        table, f"{prefix}.t{table_index}", context
                    )


def _story_for(section, story_kind: str, variant: str):
    if story_kind == "header":
        attribute = {
            "default": "header",
            "first": "first_page_header",
            "even": "even_page_header",
        }[variant]
    else:
        attribute = {
            "default": "footer",
            "first": "first_page_footer",
            "even": "even_page_footer",
        }[variant]
    return getattr(section, attribute)


def _iter_table_paragraphs(table, prefix: str, context: str):
    seen_cells = set()
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            cell_prefix = f"{prefix}.r{row_index}.c{cell_index}"
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield f"{cell_prefix}.p{paragraph_index}", paragraph, context
            for table_index, nested in enumerate(cell.tables):
                yield from _iter_table_paragraphs(
                    nested, f"{cell_prefix}.t{table_index}", context
                )


def _paragraph_source(paragraph, entry_id: str, text: str) -> tuple[str, list[str]]:
    format_info = []
    if paragraph.style and paragraph.style.name:
        format_info.append(paragraph.style.name)

    run_text = "".join(run.text for run in paragraph.runs)
    if run_text.strip() != text:
        format_info.append("包含超链接或字段；写回后转换为普通文本")
        return text, format_info

    source_parts = []
    formats_seen = set()
    for run_index, run in enumerate(paragraph.runs):
        if not run.text:
            continue
        formats = []
        if run.bold:
            formats.append(("b", "粗体", "bold"))
        if run.italic:
            formats.append(("i", "斜体", "italic"))
        if run.underline:
            formats.append(("u", "下划线", "underline"))
        for code, label, name in formats:
            formats_seen.add(name)
            marker_id = f"{entry_id}.r{run_index}.{code}"
            source_parts.append(
                f"<tag id='{marker_id}' type='fmt' desc='{label}开始'/>"
            )
        source_parts.append(run.text)
        for code, label, _name in reversed(formats):
            marker_id = f"/{entry_id}.r{run_index}.{code}"
            source_parts.append(
                f"<tag id='{marker_id}' type='/fmt' desc='{label}结束'/>"
            )
    if formats_seen:
        format_info.append(", ".join(sorted(formats_seen)))
    return "".join(source_parts) if source_parts else text, format_info


_TAG_RE = re.compile(
    r"<tag\s+id=['\"]([^'\"]+)['\"]\s+type=['\"]([^'\"]*)['\"]"
    r"\s+desc=['\"]([^'\"]*)['\"]\s*/>"
)


def _write_formatted_paragraph(paragraph, target: str) -> None:
    """Rebuild basic run formatting from protected inline markers."""
    paragraph.clear()
    active = {"bold": False, "italic": False, "underline": False}
    cursor = 0

    def add_text(text: str) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = active["bold"]
        run.italic = active["italic"]
        run.underline = active["underline"]

    for match in _TAG_RE.finditer(target):
        add_text(target[cursor:match.start()])
        tag_type = match.group(2)
        desc = match.group(3)
        enabled = not tag_type.startswith("/")
        if "粗体" in desc:
            active["bold"] = enabled
        elif "斜体" in desc:
            active["italic"] = enabled
        elif "下划线" in desc:
            active["underline"] = enabled
        cursor = match.end()
    add_text(target[cursor:])
