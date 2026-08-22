"""批处理事务、独立导出和用户源文件只读的端到端测试。"""

import io
import json
import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import batch


_MINI_MQ = """<?xml version='1.0' encoding='UTF-8'?>
<xliff xmlns="urn:oasis:names:tc:xliff:document:1.2" xmlns:mq="MQXliff" version="1.2">
<file original="t" source-language="ja" target-language="zh-cn" datatype="x-memoq">
<body>
<trans-unit id="1" mq:status="NotStarted" mq:locked="no">
<source xml:space="preserve">一</source><target xml:space="preserve"></target>
</trans-unit>
<trans-unit id="2" mq:status="NotStarted" mq:locked="no">
<source xml:space="preserve">二</source><target xml:space="preserve"></target>
</trans-unit>
</body>
</file>
</xliff>"""

_LOCKED_MQ = """<?xml version='1.0' encoding='UTF-8'?>
<xliff xmlns="urn:oasis:names:tc:xliff:document:1.2" xmlns:mq="MQXliff" version="1.2">
<file original="t" source-language="ja" target-language="zh-cn" datatype="x-memoq">
<body>
<trans-unit id="1" mq:status="NotStarted" mq:locked="locked">
<source xml:space="preserve">锁定原文</source><target xml:space="preserve"></target>
</trans-unit>
</body>
</file>
</xliff>"""


class BatchWorkflowTest(unittest.TestCase):
    def setUp(self):
        self._orig_script_dir = batch._SCRIPT_DIR
        self._orig_active = batch._ACTIVE_PROJECT
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)
        self.tool = self.base / "batch_translate"
        self.tool.mkdir()
        (self.tool / "data").mkdir()
        (self.tool / "exports").mkdir()
        shutil.copy2(ROOT / "convert.py", self.tool / "convert.py")
        shutil.copy2(ROOT / "mqxliff_tool.py", self.tool / "mqxliff_tool.py")
        shutil.copy2(ROOT / "tm_store.py", self.tool / "tm_store.py")
        shutil.copytree(ROOT / "parsers", self.tool / "parsers")
        batch._SCRIPT_DIR = self.tool
        batch._ACTIVE_PROJECT = self.tool / "data" / ".active_project"

    def tearDown(self):
        batch._SCRIPT_DIR = self._orig_script_dir
        batch._ACTIVE_PROJECT = self._orig_active
        self._tmp.cleanup()

    def _state(self, stem: str) -> dict:
        return json.loads(
            (self.tool / "data" / stem / "batch_state.json").read_text(
                encoding="utf-8"
            )
        )

    def _disable_qa(self, stem: str) -> None:
        """Keep legacy transaction tests focused on write/rollback behavior."""
        state_path = self.tool / "data" / stem / "batch_state.json"
        state = json.loads(state_path.read_text(encoding="utf-8"))
        state["qa_required"] = False
        state["qa_status"] = "not_started"
        state_path.write_text(
            json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    def _write_result(self, stem: str, batch_num: int, targets: dict[str, str]) -> Path:
        self._disable_qa(stem)
        task = json.loads(
            (self.tool / "exports" / stem / f"_batch_{batch_num:03d}_to_translate.json")
            .read_text(encoding="utf-8")
        )
        result = [
            {"id": entry["id"], "target": targets[entry["id"]]}
            for entry in task["entries"]
        ]
        path = self.base / f"result_{stem}_{batch_num}.json"
        path.write_text(json.dumps(result, ensure_ascii=False), encoding="utf-8")
        return path

    def test_txt_source_is_unchanged_and_export_is_separate(self):
        source = self.base / "dialogue.txt"
        source.write_bytes("一\n\n二\n".encode("utf-8"))
        original = source.read_bytes()

        batch.cmd_init(source, batch_chars=1)
        batch.cmd_next()
        batch.cmd_submit(self._write_result("dialogue", 1, {"1": "甲"}))
        batch.cmd_submit(self._write_result("dialogue", 2, {"3": "乙"}))

        delivered = self.base / "delivered.txt"
        batch.cmd_export("dialogue", str(delivered), False)
        self.assertEqual(source.read_bytes(), original)
        self.assertNotEqual(delivered.resolve(), source.resolve())
        self.assertEqual(delivered.read_text(encoding="utf-8"), "甲\n\n乙\n")

        with self.assertRaises(SystemExit):
            batch.cmd_export("dialogue", str(source), True)
        self.assertEqual(source.read_bytes(), original)

    def test_xlsx_custom_columns_survive_multiple_batches(self):
        source = self.base / "sheet.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Strings"
        ws.append(["meta", "unused", "source", "target"])
        ws.append(["meta", "unused", "header 2", ""])
        ws.append(["meta", "unused", "header 3", ""])
        ws.append(["m1", "NOISE1", "原文一", "既有译文"])
        ws.append(["m2", "NOISE2", "原文二", ""])
        ws.append(["m3", "NOISE3", "原文三", ""])
        wb.save(source)
        wb.close()
        original = source.read_bytes()

        batch.cmd_init(
            source,
            batch_chars=3,
            source_col="C",
            target_col="D",
            header_row=3,
            sheet_name="Strings",
        )
        batch.cmd_next()
        first_task = json.loads(
            (self.tool / "exports" / "sheet" / "_batch_001_to_translate.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual(first_task["entries"][0]["source"], "原文一")
        batch.cmd_submit(
            self._write_result("sheet", 1, {"1": "既有译文"})
        )
        second_task = json.loads(
            (self.tool / "exports" / "sheet" / "_batch_002_to_translate.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual(second_task["entries"][0]["source"], "原文二")
        batch.cmd_submit(self._write_result("sheet", 2, {"2": "译文二"}))
        third_task = json.loads(
            (self.tool / "exports" / "sheet" / "_batch_003_to_translate.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual(third_task["entries"][0]["source"], "原文三")
        batch.cmd_submit(self._write_result("sheet", 3, {"3": "译文三"}))

        delivered = self.base / "sheet_translated.xlsx"
        batch.cmd_export("sheet", str(delivered), False)
        self.assertEqual(source.read_bytes(), original)
        result_wb = openpyxl.load_workbook(delivered, data_only=True)
        result_ws = result_wb["Strings"]
        self.assertEqual(result_ws["D4"].value, "既有译文")
        self.assertEqual(result_ws["D5"].value, "译文二")
        self.assertEqual(result_ws["D6"].value, "译文三")
        result_wb.close()

    def test_docx_source_is_unchanged_and_formatting_is_exported(self):
        source = self.base / "formatted.docx"
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("原文").bold = True
        doc.save(source)
        original = source.read_bytes()

        batch.cmd_init(source)
        batch.cmd_next()
        task = json.loads(
            (self.tool / "exports" / "formatted" / "_batch_001_to_translate.json")
            .read_text(encoding="utf-8")
        )
        translated = task["entries"][0]["source"].replace("原文", "译文")
        entry_id = task["entries"][0]["id"]
        batch.cmd_submit(self._write_result("formatted", 1, {entry_id: translated}))

        delivered = self.base / "formatted_translated.docx"
        batch.cmd_export("formatted", str(delivered), False)
        self.assertEqual(source.read_bytes(), original)
        result_doc = Document(delivered)
        self.assertEqual(result_doc.paragraphs[0].text, "译文")
        self.assertTrue(result_doc.paragraphs[0].runs[0].bold)

    def test_source_locked_empty_mqxliff_completes_without_touching_source(self):
        source = self.base / "locked.mqxliff"
        source.write_text(_LOCKED_MQ, encoding="utf-8")
        original = source.read_bytes()

        batch.cmd_init(source)
        batch.cmd_next()
        task_path = self.tool / "exports" / "locked" / "_batch_001_to_translate.json"
        task = json.loads(task_path.read_text(encoding="utf-8"))
        self.assertTrue(task["entries"][0]["source_locked"])
        self.assertTrue(task["entries"][0]["locked"])
        self.assertEqual(task["entries"][0]["target"], "")
        batch.cmd_submit(self._write_result("locked", 1, {"1": ""}))

        delivered = self.base / "locked_translated.mqxliff"
        batch.cmd_export("locked", str(delivered), False)
        self.assertEqual(source.read_bytes(), original)
        self.assertTrue(delivered.is_file())

    def test_mixed_existing_target_is_never_written_or_added_to_runtime_tm(self):
        source = self.base / "mixed.mqxliff"
        source.write_text(
            _MINI_MQ.replace(
                '<target xml:space="preserve"></target>',
                '<target xml:space="preserve">既有译文</target>',
                1,
            ),
            encoding="utf-8",
        )

        batch.cmd_init(source, batch_chars=100)
        batch.cmd_next()
        task = json.loads(
            (self.tool / "exports" / "mixed" / "_batch_001_to_translate.json")
            .read_text(encoding="utf-8")
        )
        self.assertTrue(task["entries"][0]["preserve_existing"])
        self.assertTrue(task["entries"][0]["locked"])
        self._disable_qa("mixed")
        result = self.base / "mixed-result.json"
        result.write_text(
            json.dumps([
                {"id": "1", "target": "既有译文"},
                {"id": "2", "target": "新增译文"},
            ], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_submit(result)

        runtime = self.tool / "data" / "project_tm_runtime" / "mixed" / "_batch_001.json"
        runtime_entries = json.loads(runtime.read_text(encoding="utf-8"))["entries"]
        self.assertEqual(len(runtime_entries), 1)
        self.assertEqual(runtime_entries[0]["source"], "二")
        self.assertEqual(runtime_entries[0]["target"], "新增译文")
        self.assertEqual(runtime_entries[0]["file"], "_working_mixed.mqxliff")

        delivered = self.base / "mixed-delivered.mqxliff"
        batch.cmd_export("mixed", str(delivered), False)
        from mqxliff_tool import parse_mqxliff

        units, _ = parse_mqxliff(delivered)
        self.assertEqual(next(unit for unit in units if unit.id == "1").target_text, "既有译文")
        self.assertEqual(next(unit for unit in units if unit.id == "2").target_text, "新增译文")

    def test_restore_preserved_recovers_original_mqxliff_target_and_live_state(self):
        source = self.base / "restore.mqxliff"
        source.write_text(
            _MINI_MQ.replace(
                '<target xml:space="preserve"></target>',
                '<target xml:space="preserve">既有译文</target>',
                1,
            ),
            encoding="utf-8",
        )
        batch.cmd_init(source)
        state = self._state("restore")
        work_file = Path(state["source_file"])
        work_file.write_text(
            work_file.read_text(encoding="utf-8").replace("既有译文", "错误译文"),
            encoding="utf-8",
        )
        export_file = Path(state["export_file"])
        data = json.loads(export_file.read_text(encoding="utf-8"))
        data["entries"][0]["target"] = "错误译文"
        export_file.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

        batch.cmd_restore_preserved("restore")

        from mqxliff_tool import parse_mqxliff

        units, _ = parse_mqxliff(work_file)
        self.assertEqual(units[0].target_text, "既有译文")
        restored = json.loads(export_file.read_text(encoding="utf-8"))
        self.assertEqual(restored["entries"][0]["target"], "既有译文")
        live_state = self._state("restore")
        self.assertEqual(live_state["preserved_targets"]["1"], "既有译文")
        self.assertEqual(live_state["preserved_recovery"]["count"], 1)

    def test_rebuild_runtime_tm_only_includes_submitted_batches_and_updates_state(self):
        source = self.base / "runtime-rebuild.mqxliff"
        source.write_text(
            _MINI_MQ.replace(
                '<target xml:space="preserve"></target>',
                '<target xml:space="preserve">既有译文</target>',
                1,
            ),
            encoding="utf-8",
        )
        batch.cmd_init(source, batch_chars=1)
        batch.cmd_next()
        batch.cmd_submit(
            self._write_result("runtime-rebuild", 1, {"1": "既有译文"})
        )

        batch.cmd_rebuild_runtime_tm("runtime-rebuild")

        state = self._state("runtime-rebuild")
        runtime_dir = self.tool / "data" / "project_tm_runtime" / "runtime-rebuild"
        batch_1 = runtime_dir / "_batch_001.json"
        batch_2 = runtime_dir / "_batch_002.json"
        self.assertEqual(state["tm_runtime_files"], [str(batch_1.resolve())])
        self.assertFalse(batch_2.exists())
        self.assertEqual(
            json.loads(batch_1.read_text(encoding="utf-8"))["entries"], []
        )
        self.assertEqual(state["runtime_tm_rebuild"]["preserved_count"], 1)

    def test_policy_allowed_empty_clears_unlocked_mqxliff_target(self):
        source = self.base / "clear.mqxliff"
        source.write_text(
            _MINI_MQ.replace(
                '<target xml:space="preserve"></target>',
                '<target xml:space="preserve">旧译文</target>',
                1,
            ),
            encoding="utf-8",
        )
        policy = self.base / "allow-empty.json"
        policy.write_text(
            json.dumps({"allow_empty_ids": ["1"]}), encoding="utf-8"
        )
        batch.cmd_init(source, validation_policy_path=policy)
        batch.cmd_next(review_only=True)
        self._disable_qa("clear")
        result = self.base / "clear-result.json"
        result.write_text(
            json.dumps([
                {"id": "1", "target": ""},
                {"id": "2", "target": "译文二"},
            ], ensure_ascii=False),
            encoding="utf-8",
        )

        batch.cmd_submit(result)
        delivered = self.base / "clear-delivered.mqxliff"
        batch.cmd_export("clear", str(delivered), False)

        from mqxliff_tool import parse_mqxliff

        units, _ = parse_mqxliff(delivered)
        self.assertEqual(next(unit for unit in units if unit.id == "1").target_text, "")
        self.assertEqual(next(unit for unit in units if unit.id == "2").target_text, "译文二")

    def test_translate_submit_requires_review_and_qa(self):
        source = self.base / "qa_gate.txt"
        source.write_text("原文\n", encoding="utf-8")
        batch.cmd_init(source)
        batch.cmd_next()
        result = self.base / "qa-gate-result.json"
        result.write_text(
            json.dumps([{"id": "1", "target": "译文"}], ensure_ascii=False),
            encoding="utf-8",
        )

        with self.assertRaises(SystemExit):
            batch.cmd_submit(result)
        state = self._state("qa_gate")
        self.assertEqual(state["qa_status"], "awaiting_translation")
        self.assertEqual(state["current_batch"], 0)

    def test_required_translator_receipt_requires_attempt_promote_v2(self):
        source = self.base / "receipt.txt"
        source.write_text("原文\n", encoding="utf-8")
        selection = self.base / "reference-selection.json"
        selection.write_text(json.dumps({
            "approved": True,
            "style_guide": None,
            "terms": None,
            "tm_permanent": None,
            "validation_policy": None,
            "qa_policy": None,
        }), encoding="utf-8")
        role = self.base / "translator.toml"
        role.write_text(
            'name = "translator"\nmodel = "gpt-5.6-terra"\n'
            'model_reasoning_effort = "max"\n',
            encoding="utf-8",
        )

        batch.cmd_init(
            source,
            require_agent_receipts=True,
            reference_selection_path=selection,
        )
        batch.cmd_next()
        task = self.tool / "exports" / "receipt" / "_batch_001_to_translate.json"
        result = self.base / "receipt-result.json"
        result.write_text(
            json.dumps([{"id": "1", "target": "译文"}], ensure_ascii=False),
            encoding="utf-8",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_review(result, "receipt")

        wrong_role = self.base / "wrong-translator.toml"
        wrong_role.write_text(
            'name = "translator"\nmodel = "gpt-5.6-luna"\n'
            'model_reasoning_effort = "max"\n',
            encoding="utf-8",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_receipt(
                "translator", "/root/translator_001", role, task, result, "receipt"
            )

        attempt = batch.cmd_agent_attempt(
            "translator", None, "receipt", attempt_id="attempt_001",
            execution_surface="app",
        )
        attempt_task = json.loads(Path(attempt["agent_input"]).read_text(encoding="utf-8"))
        self.assertEqual(attempt_task["agent_attempt"]["outputs"]["result"], attempt["outputs"]["result"])
        self.assertEqual(attempt["agent_input_sha256"], batch._sha256_file(Path(attempt["agent_input"])))
        staged_result = Path(attempt["outputs"]["result"])
        staged_result.write_text(result.read_text(encoding="utf-8"), encoding="utf-8")
        manual_event = self.base / "translator-completed.json"
        manual_event.write_text(
            json.dumps({
                "schema_version": 1,
                "event": "agent_completed",
                "status": "completed",
                "agent_id": "/root/translator_001",
                "attempt_id": "attempt_001",
                "exit_code": 0,
                "outputs": {"result": batch._sha256_file(staged_result)},
            }),
            encoding="utf-8",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_promote(
                "translator", Path(attempt["outputs"]["result"]).parent,
                "/root/translator_001", role, manual_event, "receipt",
            )
        event = batch.cmd_agent_complete(
            "translator", Path(attempt["outputs"]["result"]).parent,
            "/root/translator_001", "receipt",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_promote(
                "translator", Path(attempt["outputs"]["result"]).parent,
                "/root/translator_001", wrong_role, event, "receipt",
            )
        batch.cmd_promote(
            "translator", Path(attempt["outputs"]["result"]).parent,
            "/root/translator_001", role, event, "receipt",
        )
        canonical_result = Path(attempt["destinations"]["result"])
        receipt = json.loads(
            (self.tool / "exports" / "receipt" / "receipts" / "_batch_001_translator.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual(receipt["schema_version"], 2)
        self.assertEqual(receipt["execution_surface"], "app")
        self.assertEqual(receipt["outputs"][0]["sha256"], batch._sha256_file(canonical_result))
        batch.cmd_review(canonical_result, "receipt")

    def test_qa_promote_is_a_two_file_receipt_transaction(self):
        source = self.base / "qa-promote.txt"
        source.write_text("こんにちは\n", encoding="utf-8")
        selection = self.base / "qa-promote-selection.json"
        selection.write_text(json.dumps({
            "approved": True,
            "style_guide": None,
            "terms": None,
            "tm_permanent": None,
            "validation_policy": None,
            "qa_policy": None,
        }), encoding="utf-8")
        translator_role = self.base / "translator-promote.toml"
        reviewer_role = self.base / "reviewer-promote.toml"
        qa_role = self.base / "qa-promote.toml"
        translator_role.write_text(
            'name = "translator"\nmodel = "gpt-5.6-terra"\nmodel_reasoning_effort = "max"\n',
            encoding="utf-8",
        )
        reviewer_role.write_text(
            'name = "trans-reviewer"\nmodel = "gpt-5.6-terra"\nmodel_reasoning_effort = "max"\n',
            encoding="utf-8",
        )
        qa_role.write_text(
            'name = "qa-reviewer"\nmodel = "gpt-5.6-luna"\nmodel_reasoning_effort = "max"\n',
            encoding="utf-8",
        )
        batch.cmd_init(source, require_agent_receipts=True, reference_selection_path=selection)
        batch.cmd_next()

        def promote(stage, role, payloads):
            attempt = batch.cmd_agent_attempt(stage, None, "qa-promote")
            for name, value in payloads.items():
                output = Path(attempt["outputs"][name])
                output.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
            event = batch.cmd_agent_complete(
                stage, Path(attempt["outputs"]["result"]).parent,
                f"/root/{stage}", "qa-promote",
            )
            batch.cmd_promote(
                stage, Path(attempt["outputs"]["result"]).parent,
                f"/root/{stage}", role, event, "qa-promote",
            )
            return attempt

        translated = promote("translator", translator_role, {
            "result": [{"id": "1", "target": "こんにちは"}],
        })
        batch.cmd_review(Path(translated["destinations"]["result"]), "qa-promote")
        reviewed = promote("trans-reviewer", reviewer_role, {
            "result": [{"id": "1", "target": "こんにちは"}],
        })
        batch.cmd_qa("qa-promote")
        qa_task = json.loads(
            (self.tool / "exports" / "qa-promote" / "_batch_001_qa_task.json").read_text(encoding="utf-8")
        )
        finding = qa_task["findings"][0]
        qa_attempt = promote("qa-reviewer", qa_role, {
            "result": [{"id": "1", "target": "你好"}],
            "report": {"schema_version": 1, "findings": [{
                "finding_id": finding["finding_id"], "id": "1", "status": "fixed", "reason": "已译为中文",
            }]},
        })
        qa_attempt_task = json.loads(
            Path(qa_attempt["agent_input"]).read_text(encoding="utf-8")
        )
        self.assertEqual(qa_attempt_task["qa_reviewed_path"], qa_attempt["outputs"]["result"])
        self.assertEqual(qa_attempt_task["qa_report_path"], qa_attempt["outputs"]["report"])
        receipt = json.loads(
            (self.tool / "exports" / "qa-promote" / "receipts" / "_batch_001_qa-reviewer.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual({item["name"] for item in receipt["outputs"]}, {"result", "report"})
        self.assertTrue(Path(qa_attempt["destinations"]["result"]).is_file())
        self.assertTrue(Path(qa_attempt["destinations"]["report"]).is_file())
        self.assertTrue(Path(reviewed["destinations"]["result"]).is_file())

    def test_project_rule_update_invalidates_all_active_documents(self):
        first = self.base / "first.txt"
        second = self.base / "second.txt"
        first.write_text("第一份\n", encoding="utf-8")
        second.write_text("第二份\n", encoding="utf-8")
        batch.cmd_init(first, project_id="document-a")
        batch.cmd_init(second, project_id="document-b")
        batch.cmd_next(project_arg="document-a")
        before = self._state("document-a")["project_rules_revision"]
        old_task = self.tool / "exports" / "document-a" / "_batch_001_to_translate.json"
        stale_result = self.base / "stale-result.json"
        stale_result.write_text(
            json.dumps([{"id": "1", "target": "第一份译文"}], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_agent_attempt("translator", None, "document-a", attempt_id="stale_attempt")
        policy = self.base / "updated-validation.json"
        policy.write_text(json.dumps({
            "newline_policy": {
                "mode": "free", "preserve_paragraphs": True, "forbid_edge_breaks": True,
            },
        }), encoding="utf-8")

        batch.cmd_project_config_update(policy, None, None, False)

        after_a = self._state("document-a")
        after_b = self._state("document-b")
        self.assertNotEqual(after_a["project_rules_revision"], before)
        self.assertEqual(after_a["project_rules_revision"], after_b["project_rules_revision"])
        self.assertEqual(after_a["qa_status"], "not_started")
        self.assertEqual(after_a["active_tasks"], {})
        self.assertEqual(after_b["active_tasks"], {})
        self.assertTrue(after_a["rules_invalidated"])
        with self.assertRaises(SystemExit):
            batch.cmd_review(stale_result, "document-a")

        batch.cmd_next(project_arg="document-a")
        regenerated = json.loads(old_task.read_text(encoding="utf-8"))
        self.assertEqual(regenerated["project_rules_revision"], after_a["project_rules_revision"])

    def test_runtime_tm_is_shared_across_document_workspaces(self):
        first = self.base / "tm-first.txt"
        second = self.base / "tm-second.txt"
        first.write_text("裏側の世界\n", encoding="utf-8")
        second.write_text("裏側の世界\n", encoding="utf-8")
        batch.cmd_init(first, project_id="tm-document-a")
        batch.cmd_next(project_arg="tm-document-a")
        result = self._write_result("tm-document-a", 1, {"1": "里侧的世界"})
        batch.cmd_submit(result, "tm-document-a")

        runtime = self.tool / "data" / "project_tm_runtime" / "tm-document-a" / "_batch_001.json"
        self.assertTrue(runtime.is_file())
        batch.cmd_init(second, project_id="tm-document-b")
        batch.cmd_next(project_arg="tm-document-b")
        task = json.loads(
            (self.tool / "exports" / "tm-document-b" / "_batch_001_to_translate.json")
            .read_text(encoding="utf-8")
        )
        self.assertEqual(task["entries"][0]["runtime_tm_matches"][0]["target"], "里侧的世界")
        self.assertEqual(task["entries"][0]["runtime_tm_matches"][0]["tm_document"], "tm-document-a")

    def test_tm_debug_reports_scored_candidates_and_rejections(self):
        source = self.base / "tm-debug.txt"
        source.write_text("裏側の世界\n", encoding="utf-8")
        permanent = self.base / "permanent.json"
        permanent.write_text(json.dumps({"entries": [
            {"source": "裏側の世界", "target": "里侧的世界", "context": "", "file": "master"},
            {"source": "ことができます", "target": "可以", "context": "", "file": "master"},
        ]}, ensure_ascii=False), encoding="utf-8")
        batch.cmd_init(source, project_id="tm-debug", tm_permanent_path=permanent)

        with mock.patch("sys.stdout", new_callable=io.StringIO) as captured:
            batch.cmd_tm_debug("裏側の世界で待つ", "tm-debug")

        report = json.loads(captured.getvalue())
        self.assertEqual(report["layers"][0]["matches"][0]["fragment_source"], "裏側の世界")
        self.assertIn("rejected", report["layers"][0])

    def test_clean_qa_rejects_changed_reviewed_baseline(self):
        source = self.base / "qa_clean.txt"
        source.write_text("原文\n", encoding="utf-8")
        batch.cmd_init(source)
        batch.cmd_next()
        translation = self.base / "qa-clean-translation.json"
        translation.write_text(
            json.dumps([{"id": "1", "target": "译文"}], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_review(translation)
        reviewed = self.tool / "exports" / "qa_clean" / "_batch_001_reviewed.json"
        reviewed.write_text(
            json.dumps([{"id": "1", "target": "译文"}], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_qa("qa_clean")

        reviewed.write_text(
            json.dumps([{"id": "1", "target": "改后"}], ensure_ascii=False),
            encoding="utf-8",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_submit(reviewed, "qa_clean")

    def test_qa_reviewer_fixes_machine_finding_before_commit(self):
        source = self.base / "qa.txt"
        source.write_text("こんにちは\n", encoding="utf-8")
        original = source.read_bytes()
        batch.cmd_init(source)
        batch.cmd_next()

        translation = self.base / "qa-translation.json"
        translation.write_text(
            json.dumps([{"id": "1", "target": "こんにちは"}], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_review(translation)

        reviewed = self.tool / "exports" / "qa" / "_batch_001_reviewed.json"
        reviewed.write_text(
            json.dumps([{"id": "1", "target": "こんにちは"}], ensure_ascii=False),
            encoding="utf-8",
        )
        batch.cmd_qa("qa")
        qa_task = json.loads(
            (self.tool / "exports" / "qa" / "_batch_001_qa_task.json").read_text(
                encoding="utf-8"
            )
        )
        finding = qa_task["findings"][0]
        qa_result = Path(qa_task["qa_reviewed_path"])
        report = Path(qa_task["qa_report_path"])
        alternate_result = self.base / "qa-reviewed.json"
        alternate_report = self.base / "qa-report.json"
        alternate_result.write_text(
            json.dumps([{"id": "1", "target": "你好"}], ensure_ascii=False),
            encoding="utf-8",
        )
        alternate_report.write_text(
            json.dumps({"schema_version": 1, "findings": []}, ensure_ascii=False),
            encoding="utf-8",
        )
        with self.assertRaises(SystemExit):
            batch.cmd_qa_submit(alternate_result, alternate_report, "qa")
        qa_result.write_text(
            json.dumps([{"id": "1", "target": "你好"}], ensure_ascii=False),
            encoding="utf-8",
        )
        report.write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "findings": [
                        {
                            "finding_id": finding["finding_id"],
                            "id": "1",
                            "status": "fixed",
                            "reason": "原文仍为日文，已改为中文译文",
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        batch.cmd_qa_submit(qa_result, report, "qa")

        manifest = json.loads(
            (self.tool / "exports" / "qa" / "project_manifest.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(manifest["qa_history"][0]["status"], "agent_reviewed")
        self.assertEqual(source.read_bytes(), original)

    def test_out_of_batch_id_is_rejected_without_file_changes(self):
        source = self.base / "ids.txt"
        source.write_text("一\n二\n", encoding="utf-8")
        batch.cmd_init(source, batch_chars=1)
        batch.cmd_next()
        self._disable_qa("ids")
        state = self._state("ids")
        tracked = [
            source,
            Path(state["source_file"]),
            Path(state["export_file"]),
            self.tool / "data" / "ids" / "batch_state.json",
        ]
        before = {path: path.read_bytes() for path in tracked}
        bad = self.base / "bad.json"
        bad.write_text(
            json.dumps([{"id": "1", "target": "甲"}, {"id": "2", "target": "乙"}], ensure_ascii=False),
            encoding="utf-8",
        )

        with self.assertRaises(SystemExit):
            batch.cmd_submit(bad)
        self.assertEqual({path: path.read_bytes() for path in tracked}, before)

    def test_invalid_validation_policy_fails_before_copying_source(self):
        source = self.base / "policy.txt"
        source.write_text("原文\n", encoding="utf-8")
        policy = self.base / "bad_policy.json"
        policy.write_text('{"unknown": true}', encoding="utf-8")

        with self.assertRaises(SystemExit):
            batch.cmd_init(source, validation_policy_path=policy)
        self.assertFalse((self.tool / "data" / "policy").exists())
        self.assertFalse(batch._ACTIVE_PROJECT.exists())

    def test_effective_validation_policy_is_injected_into_tasks(self):
        source = self.base / "policy.txt"
        source.write_text("原文\n", encoding="utf-8")
        policy = self.base / "validation_policy.json"
        policy.write_text(
            json.dumps({
                "ignored_tag_types": [],
                "allow_empty_ids": ["1"],
                "entry_overrides": {
                    "1": {
                        "tag_mode": "ignore",
                        "enforce_newline_count": True,
                    }
                },
            }),
            encoding="utf-8",
        )

        batch.cmd_init(source, validation_policy_path=policy)
        batch.cmd_next()
        task = json.loads(
            (self.tool / "exports" / "policy" / "_batch_001_to_translate.json")
            .read_text(encoding="utf-8")
        )

        validation = task["entries"][0]["validation"]
        self.assertEqual(validation["ignored_tag_types"], [])
        self.assertEqual(validation["tag_mode"], "ignore")
        self.assertTrue(validation["enforce_newline_count"])
        self.assertTrue(validation["allow_empty"])

    def test_submit_warnings_require_reason_and_are_recorded(self):
        source = self.base / "warnings.txt"
        source.write_text("原文\n", encoding="utf-8")
        batch.cmd_init(source)
        batch.cmd_next()
        self._disable_qa("warnings")
        wrapped = self.base / "wrapped.json"
        wrapped.write_text(
            json.dumps({"entries": [{"id": "1", "target": "译文"}]}, ensure_ascii=False),
            encoding="utf-8",
        )

        with self.assertRaises(SystemExit):
            batch.cmd_submit(wrapped)
        with self.assertRaises(SystemExit):
            batch.cmd_submit(wrapped, allow_warnings=True)

        batch.cmd_submit(
            wrapped,
            allow_warnings=True,
            warning_reason="项目接受对象包装",
        )

        manifest = json.loads(
            (self.tool / "exports" / "warnings" / "project_manifest.json")
            .read_text(encoding="utf-8")
        )
        acceptance = manifest["warning_acceptances"][0]
        self.assertEqual(acceptance["batch"], 1)
        self.assertEqual(acceptance["reason"], "项目接受对象包装")
        self.assertTrue(acceptance["warnings"])

    def test_same_stem_sources_use_distinct_project_ids(self):
        first_dir = self.base / "project_a"
        second_dir = self.base / "project_b"
        first_dir.mkdir()
        second_dir.mkdir()
        first = first_dir / "dialogue.txt"
        second = second_dir / "dialogue.txt"
        first.write_text("一\n", encoding="utf-8")
        second.write_text("二\n", encoding="utf-8")

        first_project = batch.cmd_init(first)
        second_project = batch.cmd_init(second)

        self.assertEqual(first_project, "dialogue")
        self.assertNotEqual(second_project, first_project)
        self.assertTrue(second_project.startswith("dialogue-"))
        first_state = self._state(first_project)
        second_state = self._state(second_project)
        self.assertEqual(Path(first_state["input_source_file"]), first.resolve())
        self.assertEqual(Path(second_state["input_source_file"]), second.resolve())
        self.assertEqual(
            batch._ACTIVE_PROJECT.read_text(encoding="utf-8"), second_project
        )

    def test_context_split_and_pack_cover_all_entries(self):
        source = self.base / "large.txt"
        source.write_text("一一\n二二\n三三\n", encoding="utf-8")
        batch.cmd_init(source)

        manifest_path = batch.cmd_context_split(3, "large")
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

        self.assertEqual(manifest["total_parts"], 3)
        part_entries = []
        report_paths = []
        for part in manifest["parts"]:
            payload = json.loads(Path(part["path"]).read_text(encoding="utf-8"))
            part_entries.extend(entry["id"] for entry in payload["entries"])
            report = self.base / f"part_{part['part']}.md"
            report.write_text(f"分片 {part['part']} 报告", encoding="utf-8")
            report_paths.append(report)
        self.assertEqual(part_entries, ["1", "2", "3"])

        merge_path = batch.cmd_context_pack(report_paths, None, "large")
        merge = json.loads(merge_path.read_text(encoding="utf-8"))

        self.assertEqual(merge["mode"], "context_merge")
        self.assertEqual(len(merge["part_reports"]), 3)

    def test_context_split_projects_only_analysis_fields(self):
        source = self.base / "context-fields.txt"
        source.write_text("一\n二\n", encoding="utf-8")
        batch.cmd_init(source)
        export_file = self.tool / "exports" / "context-fields" / "_working.json"
        data = json.loads(export_file.read_text(encoding="utf-8"))
        data["entries"][0]["tm_matches"] = [{"source": "噪声", "target": "noise"}]
        data["entries"][0]["terms"] = [{"source": "噪声", "target": "noise"}]
        export_file.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

        manifest_path = batch.cmd_context_split(100, "context-fields")
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        part = json.loads(Path(manifest["parts"][0]["path"]).read_text(encoding="utf-8"))
        self.assertNotIn("tm_matches", part["entries"][0])
        self.assertNotIn("terms", part["entries"][0])
        self.assertEqual(part["entries"][0]["source"], "一")

    def test_init_refuses_to_overwrite_existing_state(self):
        source = self.base / "existing.txt"
        source.write_text("一\n", encoding="utf-8")
        project_id = batch.cmd_init(source)
        state_path = self.tool / "data" / project_id / "batch_state.json"
        before = state_path.read_bytes()

        with self.assertRaises(SystemExit):
            batch.cmd_init(source)

        self.assertEqual(state_path.read_bytes(), before)

    def test_mqxliff_failure_rolls_back_work_json_tm_and_state(self):
        source = self.base / "rollback.mqxliff"
        source.write_text(_MINI_MQ, encoding="utf-8")
        tm = self.base / "tm.json"
        tm.write_text('{"entries": []}', encoding="utf-8")
        batch.cmd_init(source, batch_chars=1, tm_path=tm)
        batch.cmd_next()
        self._disable_qa("rollback")
        state = self._state("rollback")
        tracked = [
            source,
            Path(state["source_file"]),
            Path(state["export_file"]),
            tm,
            self.tool / "data" / "rollback" / "batch_state.json",
        ]
        before = {path: path.read_bytes() for path in tracked}
        result = self._write_result("rollback", 1, {"1": "甲"})
        real_run = __import__("subprocess").run

        def fail_reparse(args, *pargs, **kwargs):
            if "convert.py" in str(args[1]) and "parse" in args:
                raise __import__("subprocess").CalledProcessError(1, args)
            return real_run(args, *pargs, **kwargs)

        with mock.patch("subprocess.run", side_effect=fail_reparse):
            with self.assertRaises(__import__("subprocess").CalledProcessError):
                batch.cmd_submit(result)
        self.assertEqual({path: path.read_bytes() for path in tracked}, before)

    def test_layered_tm_keeps_permanent_read_only_and_runtime_per_batch(self):
        source = self.base / "layered.mqxliff"
        source.write_text(_MINI_MQ.replace("二", "一"), encoding="utf-8")
        permanent = self.base / "permanent.json"
        permanent.write_text(
            json.dumps({"entries": [{
                "source": "一", "target": "权威", "context": "", "file": "master"
            }]}, ensure_ascii=False),
            encoding="utf-8",
        )
        permanent_before = permanent.read_bytes()

        project_id = batch.cmd_init(
            source,
            batch_chars=1,
            tm_permanent_path=permanent,
        )
        batch.cmd_next()
        first = self._write_result(project_id, 1, {"1": "甲"})
        self._disable_qa(project_id)
        batch.cmd_submit(first, project_id)

        runtime_dir = self.tool / "data" / "project_tm_runtime" / project_id
        runtime_1 = runtime_dir / "_batch_001.json"
        self.assertTrue(runtime_1.is_file())
        self.assertEqual(permanent.read_bytes(), permanent_before)
        self.assertEqual(json.loads(runtime_1.read_text(encoding="utf-8"))["entries"][0]["target"], "甲")

        next_task = self.tool / "exports" / project_id / "_batch_002_to_translate.json"
        next_data = json.loads(next_task.read_text(encoding="utf-8"))
        self.assertEqual(next_data["entries"][0]["runtime_tm_matches"][0]["target"], "甲")

        second = self._write_result(project_id, 2, {"2": "乙"})
        batch.cmd_submit(second, project_id)
        runtime_2 = runtime_dir / "_batch_002.json"
        self.assertTrue(runtime_2.is_file())
        self.assertEqual(permanent.read_bytes(), permanent_before)
        self.assertEqual(json.loads(runtime_2.read_text(encoding="utf-8"))["entries"][0]["target"], "乙")

    def test_init_rejects_permanent_tm_inside_runtime_directory(self):
        source = self.base / "unsafe.mqxliff"
        source.write_text(_MINI_MQ, encoding="utf-8")
        runtime_dir = self.base / "runtime"
        permanent = runtime_dir / "permanent.json"
        with self.assertRaises(SystemExit):
            batch.cmd_init(
                source,
                tm_permanent_path=permanent,
                tm_runtime_dir=runtime_dir,
            )
        self.assertFalse((self.tool / "data" / "unsafe").exists())

    def test_state_write_failure_rolls_back_everything(self):
        source = self.base / "state.txt"
        source.write_text("一\n二\n", encoding="utf-8")
        batch.cmd_init(source, batch_chars=1)
        batch.cmd_next()
        self._disable_qa("state")
        state = self._state("state")
        state_path = self.tool / "data" / "state" / "batch_state.json"
        tracked = [source, Path(state["source_file"]), Path(state["export_file"]), state_path]
        before = {path: path.read_bytes() for path in tracked}
        result = self._write_result("state", 1, {"1": "甲"})

        def corrupt_then_fail(updated_state):
            state_path.write_text('{"corrupt": true}', encoding="utf-8")
            raise OSError("injected state write failure")

        with mock.patch.object(batch, "_save_state", side_effect=corrupt_then_fail):
            with self.assertRaises(OSError):
                batch.cmd_submit(result)
        self.assertEqual({path: path.read_bytes() for path in tracked}, before)


if __name__ == "__main__":
    unittest.main()
