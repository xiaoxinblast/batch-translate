"""txt / xlsx 解析与写回回归测试。"""

import json
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from parsers import docx_parser, txt_parser, xlsx_parser


class TxtParserTest(unittest.TestCase):
    def test_write_keeps_line_number_ids_with_blank_lines(self):
        """带空行的 txt：id 为行号（空行占号），写回按行号定位、空行原样保留。"""
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            src = td / "sample.txt"
            src.write_text("こんにちは\n\nさようなら\n", encoding="utf-8")

            data = txt_parser.parse(src)
            self.assertEqual([e["id"] for e in data["entries"]], ["1", "3"])
            data["entries"][0]["target"] = "你好"
            data["entries"][1]["target"] = "再见"  # id 3（行号）

            export = td / "parsed.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            out = txt_parser.write(src, export)

            self.assertEqual(out.read_text(encoding="utf-8"), "你好\n\n再见\n")

    def test_write_plain_array_fallback(self):
        """纯数组输入仍按行序写回。"""
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            src = td / "sample.txt"
            src.write_text("A\nB\n", encoding="utf-8")
            trans = td / "trans.json"
            trans.write_text(
                json.dumps([{"id": "1", "target": "甲"}, {"id": "2", "target": "乙"}], ensure_ascii=False),
                encoding="utf-8",
            )
            out = txt_parser.write(src, trans)
            self.assertEqual(out.read_text(encoding="utf-8"), "甲\n乙\n")

    def test_write_can_clear_existing_text(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            src = td / "sample.txt"
            src.write_text("旧译文\n", encoding="utf-8")
            data = txt_parser.parse(src)
            data["entries"][0]["target"] = ""
            trans = td / "trans.json"
            trans.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            out = txt_parser.write(src, trans)

            self.assertEqual(out.read_text(encoding="utf-8"), "\n")


class XlsxParserTest(unittest.TestCase):
    def _make_wb(self, path: Path):
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="说明行")
        ws.cell(row=2, column=1, value="说明行2")
        ws.cell(row=3, column=1, value="原文")
        ws.cell(row=3, column=2, value="注释")
        ws.cell(row=3, column=3, value="译文")
        ws.cell(row=4, column=1, value="セフィロス")
        ws.cell(row=4, column=2, value="FF7角色")
        ws.cell(row=4, column=3, value="旧译文")
        ws.cell(row=5, column=1, value="ティファ")
        ws.cell(row=5, column=2, value="FF7角色")
        wb.save(str(path))

    def test_parse_skips_header_rows(self):
        """表头行本身不进入条目，id 从表头下一行开始为 1。"""
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "sample.xlsx"
            self._make_wb(src)
            data = xlsx_parser.parse(src, source_col="A", target_col="C", header_row=3)
            self.assertEqual([e["id"] for e in data["entries"]], ["1", "2"])
            self.assertEqual(data["entries"][0]["source"], "セフィロス")
            self.assertEqual(data["entries"][0]["target"], "旧译文")
            self.assertEqual(data["entries"][0]["_row"], 4)
            self.assertEqual(data["entries"][0]["_target_col"], 2)

    def test_parse_header_row_zero_includes_first_row(self):
        """--header-row 0 表示无表头，数据从第 1 行开始。"""
        with tempfile.TemporaryDirectory() as td:
            import openpyxl

            src = Path(td) / "sample.xlsx"
            wb = openpyxl.Workbook()
            wb.active.cell(row=1, column=1, value="セフィロス")
            wb.save(str(src))
            data = xlsx_parser.parse(src, header_row=0)
            self.assertEqual([e["id"] for e in data["entries"]], ["1"])
            self.assertEqual(data["entries"][0]["_row"], 1)

    def test_write_uses_row_and_target_col(self):
        """写回时按 _row / _target_col 定位，自定义 C 列生效。"""
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "sample.xlsx"
            self._make_wb(src)
            data = xlsx_parser.parse(src, source_col="A", target_col="C", header_row=3)
            data["entries"][0]["target"] = "萨菲罗斯"
            data["entries"][1]["target"] = "蒂法"

            export = Path(td) / "parsed.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            out = xlsx_parser.write(src, export)

            import openpyxl

            ws = openpyxl.load_workbook(out).active
            self.assertEqual(ws.cell(row=4, column=3).value, "萨菲罗斯")
            self.assertEqual(ws.cell(row=5, column=3).value, "蒂法")

    def test_write_can_clear_existing_target_cell(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "sample.xlsx"
            self._make_wb(src)
            data = xlsx_parser.parse(src, source_col="A", target_col="C", header_row=3)
            data["entries"][0]["target"] = ""
            export = Path(td) / "parsed.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            out = xlsx_parser.write(src, export)

            import openpyxl

            ws = openpyxl.load_workbook(out).active
            self.assertIsNone(ws.cell(row=4, column=3).value)

    def test_source_column_to_the_right_is_parsed(self):
        with tempfile.TemporaryDirectory() as td:
            import openpyxl

            src = Path(td) / "rightmost.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["译文", "", "", "原文"])
            ws.append(["已有译文", "", "", "日本語"])
            wb.save(src)
            data = xlsx_parser.parse(
                src, source_col="D", target_col="A", header_row=1
            )
            self.assertEqual(len(data["entries"]), 1)
            self.assertEqual(data["entries"][0]["source"], "日本語")
            self.assertEqual(data["entries"][0]["target"], "已有译文")

    def test_all_sheets_roundtrip(self):
        with tempfile.TemporaryDirectory() as td:
            import openpyxl

            src = Path(td) / "multi.xlsx"
            wb = openpyxl.Workbook()
            wb.active.title = "One"
            wb.active.append(["原文", "译文"])
            wb.active.append(["一", ""])
            ws2 = wb.create_sheet("Two")
            ws2.append(["原文", "译文"])
            ws2.append(["二", ""])
            wb.save(src)

            data = xlsx_parser.parse(src, sheet_name="*")
            self.assertEqual([e["_sheet"] for e in data["entries"]], ["One", "Two"])
            data["entries"][0]["target"] = "甲"
            data["entries"][1]["target"] = "乙"
            export = Path(td) / "multi.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            out = xlsx_parser.write(src, export)
            result = openpyxl.load_workbook(out, data_only=True)
            self.assertEqual(result["One"]["B2"].value, "甲")
            self.assertEqual(result["Two"]["B2"].value, "乙")
            result.close()


class DocxParserTest(unittest.TestCase):
    def test_write_restores_basic_run_formatting(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "format.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("粗体").bold = True
            paragraph.add_run("普通")
            doc.save(src)

            data = docx_parser.parse(src)
            source = data["entries"][0]["source"]
            self.assertIn("粗体开始", source)
            data["entries"][0]["target"] = source.replace("/>粗体<tag", "/>加粗<tag")
            export = Path(td) / "format.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            out = docx_parser.write(src, export)

            result = Document(out)
            runs = result.paragraphs[0].runs
            self.assertEqual("".join(run.text for run in runs), "加粗普通")
            self.assertTrue(runs[0].bold)
            self.assertFalse(bool(runs[-1].bold))

    def test_write_can_clear_existing_paragraph(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "clear.docx"
            doc = Document()
            doc.add_paragraph("旧译文")
            doc.save(src)
            data = docx_parser.parse(src)
            data["entries"][0]["target"] = ""
            export = Path(td) / "clear.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            out = docx_parser.write(src, export)

            self.assertEqual(Document(out).paragraphs[0].text, "")

    def test_position_ids_do_not_shift_after_clearing_paragraph(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "stable.docx"
            doc = Document()
            doc.add_paragraph("第一段")
            doc.add_paragraph("")
            doc.add_paragraph("第三段")
            doc.save(src)
            data = docx_parser.parse(src)
            ids = [entry["id"] for entry in data["entries"]]
            self.assertEqual(ids, ["body.p0", "body.p2"])
            data["entries"][0]["target"] = ""
            data["entries"][1]["target"] = "第三段译文"
            export = Path(td) / "stable.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            out = docx_parser.write(src, export)
            reparsed = docx_parser.parse(out)

            self.assertEqual(reparsed["entries"][0]["id"], "body.p2")
            self.assertEqual(reparsed["entries"][0]["source"], "第三段译文")

    def test_all_table_cells_nested_tables_headers_and_footers_roundtrip(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "stories.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "左单元格"
            table.cell(0, 1).text = "右单元格"
            nested = table.cell(0, 1).add_table(rows=1, cols=1)
            nested.cell(0, 0).text = "嵌套单元格"
            doc.sections[0].header.paragraphs[0].text = "页眉"
            doc.sections[0].footer.paragraphs[0].text = "页脚"
            doc.save(src)

            data = docx_parser.parse(src)
            by_source = {entry["source"]: entry for entry in data["entries"]}
            expected = {"左单元格", "右单元格", "嵌套单元格", "页眉", "页脚"}
            self.assertTrue(expected.issubset(by_source))
            for source in expected:
                by_source[source]["target"] = f"译-{source}"
            export = Path(td) / "stories.json"
            export.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

            out = docx_parser.write(src, export)
            reparsed = docx_parser.parse(out)
            translated = {entry["source"] for entry in reparsed["entries"]}

            self.assertTrue({f"译-{source}" for source in expected}.issubset(translated))

    def test_text_box_is_reported_as_unsupported(self):
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "textbox.docx"
            doc = Document()
            run = doc.add_paragraph("正文").add_run()
            run._r.append(parse_xml(
                f'<w:pict {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml">'
                '<v:shape><v:textbox><w:txbxContent><w:p><w:r>'
                '<w:t>文本框内容</w:t></w:r></w:p></w:txbxContent>'
                '</v:textbox></v:shape></w:pict>'
            ))
            doc.save(src)

            data = docx_parser.parse(src)

            self.assertTrue(any("文本框" in warning for warning in data["warnings"]))


if __name__ == "__main__":
    unittest.main()
